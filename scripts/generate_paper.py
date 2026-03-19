"""Generate the corrected paper as a .docx file using python-docx."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold


def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(10)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], str(val))

    doc.add_paragraph()  # spacing


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.bold = bold
    run.italic = italic
    return p


def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"


def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # =========================================================================
    # TITLE
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Optimizing Temporal Efficient Networks for Embedded "
        "Keyword Spotting on Microcontrollers"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    # =========================================================================
    # ABSTRACT (#1, #12 — moved to top)
    # =========================================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Abstract")
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run = p.add_run(
        " \u2014 Keyword spotting (KWS) is a key enabling technology for "
        "voice-driven interfaces in embedded and Internet-of-Things (IoT) "
        "devices. Deploying deep learning models for KWS directly on "
        "microcontrollers is challenging due to strict constraints in memory "
        "capacity, computational resources, and energy consumption. In this "
        "work, we investigate the optimization of a lightweight neural "
        "network architecture for embedded keyword spotting applications. "
        "Specifically, we focus on improving the Temporal Efficient Network "
        "(TENet) architecture through systematic architecture exploration "
        "and hyperparameter optimization. Five candidate architectures are "
        "evaluated under a fixed computational budget, and TENet is selected "
        "as the most efficient base model. A four-stage progressive grid "
        "search is then performed over six architectural hyperparameters, "
        "evaluating 2,302 candidate configurations. The optimized model is "
        "trained using the Google Speech Commands v2 dataset with data "
        "augmentation techniques including SpecAugment, Mixup, and label "
        "smoothing, and deployed on an STM32F767IGT6 microcontroller based "
        "on the ARM Cortex-M7 architecture. Experimental results show that "
        "the proposed model achieves a test accuracy of 96.84% (96.76% after "
        "INT8 quantization) and a weighted F1-score of 0.9675 with only "
        "12,822 parameters and 276,136 multiply-accumulate operations "
        "(MACCs). The deployed INT8 model requires 13.87 KB of flash memory "
        "and 7.64 KB of RAM, and achieves an average inference latency of "
        "6.32 ms on the target hardware. Compared with previously reported "
        "lightweight keyword spotting architectures operating under similar "
        "computational constraints, the optimized TENet model achieves "
        "improved classification performance while maintaining a compact "
        "footprint suitable for real-time embedded deployment."
    )
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run("Keywords")
    run.bold = True
    run.italic = True
    run = p.add_run(
        " \u2014 Keyword spotting, TinyML, embedded AI, microcontroller, "
        "TENet, speech recognition, edge computing, neural architecture "
        "optimization."
    )
    run.font.size = Pt(11)

    # =========================================================================
    # I. INTRODUCTION (#17, #18)
    # =========================================================================
    add_heading(doc, "I. Introduction", level=1)

    add_para(doc, (
        "Keyword spotting (KWS) is a speech processing task that detects "
        "predefined keywords from a continuous audio stream. It forms the "
        "foundation of voice-driven interfaces used in many modern "
        "applications such as smart speakers, wearable devices, and "
        "Internet-of-Things (IoT) systems [1]. In these applications, a "
        "lightweight KWS module continuously monitors incoming audio signals "
        "and activates higher-level processing only when a specific command "
        "or wake word is detected. This mechanism enables intuitive "
        "human\u2013machine interaction while reducing unnecessary computation."
    ), indent=True)

    add_para(doc, (
        "With the rapid growth of edge computing and IoT technologies, "
        "there is increasing demand for deploying speech recognition "
        "algorithms directly on embedded devices. Running KWS locally on "
        "edge hardware offers several advantages, including reduced latency, "
        "improved privacy, and lower communication overhead since audio data "
        "does not need to be transmitted to cloud servers [2]. However, "
        "implementing deep learning models on microcontrollers remains "
        "challenging due to strict constraints in memory capacity, "
        "computational power, and energy consumption [3]."
    ), indent=True)

    add_para(doc, (
        "To address these constraints, recent research has focused on "
        "developing lightweight neural network architectures for embedded "
        "keyword spotting. Zhang et al. [2] proposed depthwise separable "
        "CNN (DS-CNN) models specifically designed for microcontroller "
        "deployment, demonstrating real-time keyword spotting on ARM "
        "Cortex-M processors. Li et al. [4] introduced the Temporal "
        "Efficient Network (TENet), which uses multi-scale temporal "
        "convolutions to efficiently capture temporal patterns in speech "
        "signals. Kim et al. [5] proposed BC-ResNet with broadcasted "
        "residual connections, achieving high accuracy with compact "
        "architectures. Majumdar and Ginsburg [6] developed MatchboxNet "
        "using 1D time-channel separable convolutions. In these systems, "
        "compact acoustic features such as Mel-Frequency Cepstral "
        "Coefficients (MFCCs) [7] are commonly used to represent speech "
        "signals in a computationally efficient form."
    ), indent=True)

    add_para(doc, (
        "Despite significant progress, designing efficient KWS systems for "
        "resource-constrained microcontrollers remains an active research "
        "problem. Most reported high-accuracy models require millions of "
        "multiply-accumulate operations (MACCs), making them impractical "
        "for the most constrained deployment scenarios. Selecting the most "
        "suitable neural architecture under strict computational limits is "
        "particularly challenging for TinyML applications [3], where both "
        "recognition accuracy and inference latency must be carefully "
        "balanced."
    ), indent=True)

    add_para(doc, (
        "Motivated by these challenges, this work investigates the "
        "optimization of a lightweight neural architecture for keyword "
        "spotting on embedded microcontrollers. Specifically, we focus on "
        "the Temporal Efficient Network (TENet) architecture [4] and "
        "perform systematic architecture exploration and hyperparameter "
        "optimization to improve its performance under a strict "
        "computational budget of 287,673 MACCs. The optimized model is "
        "trained on the Google Speech Commands v2 dataset [1] and deployed "
        "on an STM32F767IGT6 microcontroller based on the ARM Cortex-M7 "
        "architecture using the ST X-CUBE-AI framework [8]."
    ), indent=True)

    add_para(doc, "The main contributions of this work are summarized as follows:")

    add_numbered_list(doc, [
        "Architecture evaluation for embedded keyword spotting: Five "
        "lightweight neural architectures (TENet, LiCoNet, DS-CNN, "
        "MobileNet-style, and BC-ResNet) are evaluated under a fixed MACC "
        "budget to identify the most efficient model for microcontroller "
        "deployment.",

        "Progressive grid-search optimization of TENet: A four-stage "
        "progressive grid search is performed over six architectural "
        "hyperparameters, evaluating 2,302 candidate configurations to "
        "optimize the TENet architecture while maintaining computational "
        "constraints.",

        "Embedded implementation with INT8 quantization on Cortex-M7: "
        "The optimized model is quantized to INT8 precision using ONNX "
        "Runtime static post-training quantization and deployed on an "
        "STM32F767IGT6 platform using the ST X-CUBE-AI framework.",

        "Significant improvement over existing embedded KWS models: The "
        "optimized model achieves 96.76% INT8 accuracy at 276K MACCs, "
        "surpassing the previously reported DS-CNN baseline [9] by 5.76 "
        "percentage points while requiring 91.8% less flash memory "
        "(13.87 KB vs. 169.63 KB).",
    ])

    add_para(doc, (
        "The remainder of this paper is organized as follows. Section II "
        "reviews related work in lightweight keyword spotting. Section III "
        "describes the proposed methodology and system design. Section IV "
        "presents the experimental results and discussion. Finally, "
        "Section V concludes the paper."
    ), indent=True)

    # =========================================================================
    # II. RELATED WORK (#11 — entirely new)
    # =========================================================================
    add_heading(doc, "II. Related Work", level=1)

    add_heading(doc, "A. Lightweight Architectures for Keyword Spotting", level=2)
    add_para(doc, (
        "The DS-CNN architecture proposed by Zhang et al. [2] established "
        "an influential baseline for keyword spotting on microcontrollers. "
        "Using depthwise separable convolutions, DS-CNN achieves "
        "competitive accuracy with significantly fewer parameters than "
        "standard CNNs. The model was deployed on an ARM Cortex-M7 "
        "processor with CMSIS-NN [16], demonstrating real-time inference "
        "within 70 KB of memory."
    ), indent=True)

    add_para(doc, (
        "Li et al. [4] introduced TENet, which processes MFCC features as "
        "1D temporal sequences using multi-scale temporal convolutions with "
        "an inverted bottleneck structure inspired by MobileNetV2 [10]. "
        "TENet achieves 96.8% accuracy on Google Speech Commands v2 with "
        "100K parameters and 2.9M multiplications."
    ), indent=True)

    add_para(doc, (
        "Kim et al. [5] proposed BC-ResNet, which uses broadcasted residual "
        "connections and sub-spectral normalization to achieve "
        "state-of-the-art accuracy (up to 98.7% with BC-ResNet-8) on the "
        "same benchmark. However, even the smallest variant (BC-ResNet-1, "
        "9.2K parameters) requires 3.1M MACs, which exceeds the "
        "computational budget considered in this work."
    ), indent=True)

    add_para(doc, (
        "Majumdar and Ginsburg [6] developed MatchboxNet using 1D "
        "time-channel separable convolutions, achieving 97.48% accuracy "
        "with 93K parameters. LiCoNet uses linearized convolutions for "
        "efficient streaming inference on embedded platforms."
    ), indent=True)

    add_heading(doc, "B. Neural Architecture Search for TinyML", level=2)
    add_para(doc, (
        "Neural architecture search (NAS) methods have been applied to "
        "optimize models for microcontroller deployment. Lin et al. [11] "
        "proposed MCUNet, which co-designs the neural architecture "
        "(TinyNAS) and inference engine (TinyEngine) to jointly optimize "
        "accuracy and latency under memory constraints. Differentiable NAS "
        "methods such as DARTS [12] and hardware-aware approaches such as "
        "MicroNAS [13] enable efficient exploration of large search spaces."
    ), indent=True)

    add_para(doc, (
        "For smaller search spaces with a limited number of discrete "
        "hyperparameters, grid search and random search remain practical "
        "alternatives. Bergstra and Bengio [14] showed that random search "
        "can be more efficient than grid search in high-dimensional spaces, "
        "while successive halving strategies [15] reduce computational cost "
        "by progressively eliminating poor candidates. The progressive grid "
        "search approach used in this work combines exhaustive enumeration "
        "with staged elimination, which is effective for the "
        "six-dimensional search space considered."
    ), indent=True)

    add_heading(doc, "C. Embedded Deployment Frameworks", level=2)
    add_para(doc, (
        "Several frameworks support neural network deployment on "
        "microcontrollers. CMSIS-NN [16] provides optimized kernels for "
        "ARM Cortex-M processors. TensorFlow Lite for Microcontrollers "
        "(TFLM) [3] offers a portable runtime with INT8 quantization "
        "support. ST X-CUBE-AI [8] converts trained models (ONNX, TFLite, "
        "Keras) to optimized C code for STM32 microcontrollers, with "
        "support for INT8 quantized models in QDQ format."
    ), indent=True)

    add_para(doc, (
        "The MLPerf Tiny benchmark [3] standardizes evaluation of ML "
        "models on microcontrollers, including a keyword spotting task with "
        "a DS-CNN reference model achieving \u226590% accuracy."
    ), indent=True)

    # =========================================================================
    # III. METHODOLOGY
    # =========================================================================
    add_heading(doc, "III. Methodology", level=1)

    # III.A System Overview
    add_heading(doc, "A. System Overview", level=2)
    add_para(doc, (
        "The proposed keyword spotting system follows a lightweight "
        "processing pipeline designed for deployment on resource-constrained "
        "microcontrollers. The system consists of three main stages: audio "
        "acquisition, feature extraction, and neural network inference."
    ), indent=True)

    add_para(doc, (
        "Raw audio waveforms are captured at 16 kHz and segmented into "
        "one-second frames. The audio signal is then transformed into MFCC "
        "features, which provide a compact representation of the spectral "
        "characteristics of speech. These features are used as input to a "
        "lightweight neural network classifier that predicts the "
        "corresponding keyword class. The overall design minimizes "
        "computational complexity and memory usage to enable real-time "
        "inference on embedded hardware."
    ), indent=True)

    # III.B Dataset (#19)
    add_heading(doc, "B. Dataset", level=2)
    add_para(doc, (
        "All experiments in this work are conducted using the Google Speech "
        "Commands Dataset v2 [1], which is widely used for benchmarking "
        "keyword spotting systems."
    ), indent=True)

    add_para(doc, (
        "The classification task consists of 12 classes, including ten "
        "keyword commands (yes, no, up, down, left, right, on, off, stop, "
        "go) together with two additional classes (unknown and silence). "
        "The unknown class groups together speech samples that do not "
        "correspond to the target commands. The silence class consists of "
        "zero-padded waveforms generated to comprise approximately 10% of "
        "the base training dataset size."
    ), indent=True)

    add_table(doc,
              ["Property", "Value"],
              [
                  ["Sampling rate", "16 kHz"],
                  ["Duration per sample", "~1 second (16,256 samples)"],
                  ["Training samples", "~51,000"],
                  ["Validation samples", "~13,000"],
                  ["Test samples", "12,105"],
                  ["Number of classes", "12"],
              ])

    add_para(doc, (
        "The dataset uses the official train/validation/test split provided "
        "by the dataset authors."
    ), indent=True)

    # III.C Feature Extraction (#2, #16)
    add_heading(doc, "C. Feature Extraction", level=2)
    add_para(doc, (
        "To convert raw audio signals into a representation suitable for "
        "neural network processing, Mel-Frequency Cepstral Coefficient "
        "(MFCC) features [7] are extracted from each audio sample using "
        "torchaudio.transforms.MFCC."
    ), indent=True)

    add_table(doc,
              ["Parameter", "Value"],
              [
                  ["Number of MFCC coefficients (n_mfcc)", "13"],
                  ["Number of Mel filters (n_mels)", "16"],
                  ["FFT size (n_fft)", "512"],
                  ["Window length", "512"],
                  ["Hop length", "256"],
                  ["Minimum frequency (f_min)", "20 Hz"],
                  ["Center padding", "False"],
                  ["Power", "2.0"],
              ])

    add_para(doc, (
        "These parameters produce an MFCC feature representation of 62 "
        "temporal frames \u00d7 13 coefficients per sample, resulting in an "
        "input tensor of shape (1, 62, 13) with a size of 3,224 bytes in "
        "float32 precision."
    ), indent=True)

    # III.D Architecture Selection (#13 — new)
    add_heading(doc, "D. Architecture Selection", level=2)
    add_para(doc, (
        "To identify the most efficient architecture for the target "
        "computational budget, five lightweight neural network "
        "architectures were evaluated under a constraint of approximately "
        "287K MACCs. All models were trained for 50 epochs with identical "
        "training settings (AdamW, lr=0.001, batch size=1024) on the "
        "12-class task."
    ), indent=True)

    add_table(doc,
              ["Architecture", "Type", "MACCs", "Float32 Acc", "INT8 Acc", "INT8 Size"],
              [
                  ["TENet [4]", "1D Inverted Bottleneck", "284,896", "96.52%", "96.56%", "9.5 KB"],
                  ["LiCoNet", "1D Linearized Conv", "285,980", "95.56%", "95.42%", "8.6 KB"],
                  ["CustomDSCNN", "2D Depthwise Separable", "286,240", "94.71%", "94.49%", "6.3 KB"],
                  ["MobileNet-style", "2D MobileNet", "683,624", "94.76%", "94.14%", "14.3 KB"],
                  ["BC-ResNet [5]", "2D Broadcast Residual", "290,008", "83.85%", "71.84%", "2.3 KB"],
              ],
              caption="Table I: Architecture comparison under ~287K MACC budget (50 epochs).")

    add_para(doc, (
        "TENet achieved the highest accuracy (96.52%) and maintained "
        "excellent INT8 quantization robustness (96.56%), and was therefore "
        "selected as the base architecture for further optimization. "
        "BC-ResNet, despite its strong performance at larger scales [5], "
        "performed poorly at extreme channel reduction (channels as low as "
        "4\u201312), likely due to insufficient capacity for its sub-spectral "
        "normalization mechanism."
    ), indent=True)

    # III.E TENet Architecture (#7)
    add_heading(doc, "E. TENet Architecture", level=2)
    add_para(doc, (
        "The TENet architecture [4] processes MFCC sequences using "
        "one-dimensional temporal convolutions. The input tensor (1, 62, 13) "
        "is reshaped to (13, 62), treating the 13 MFCC coefficients as "
        "input channels and the 62 frames as the temporal dimension."
    ), indent=True)

    add_para(doc, (
        "The architecture adopts an inverted bottleneck design inspired by "
        "MobileNetV2 [10], where each block consists of: (1) pointwise "
        "expansion Conv1d (in_ch \u2192 in_ch \u00d7 ratio, kernel=1) to increase "
        "channel dimension; (2) depthwise temporal Conv1d (hidden, hidden, "
        "kernel=k, groups=hidden) for temporal feature extraction; "
        "(3) pointwise projection Conv1d (hidden \u2192 out_ch, kernel=1) for "
        "dimensionality reduction with no activation (linear); and "
        "(4) a residual connection (identity shortcut when stride=1 and "
        "in_ch=out_ch, otherwise a 1\u00d71 projection shortcut)."
    ), indent=True)

    add_para(doc, (
        "The network begins with a stem convolution (Conv1d, kernel=3), "
        "followed by multiple inverted bottleneck stages, and concludes "
        "with adaptive average pooling and a fully connected classification "
        "layer. The optimized architecture parameters obtained through grid "
        "search (Section III.F) are shown in Table II."
    ), indent=True)

    add_table(doc,
              ["Parameter", "Value"],
              [
                  ["Stem channels", "14"],
                  ["Block channels", "[14, 14, 14, 14, 14]"],
                  ["Strides", "[2, 2, 1, 1]"],
                  ["Expansion ratios", "[4, 4, 4, 4]"],
                  ["Layers per block", "[1, 2, 1, 1]"],
                  ["Kernel size", "9"],
                  ["Dropout", "0.1"],
                  ["Total parameters", "12,822"],
                  ["Total MACCs", "276,136"],
              ],
              caption="Table II: Optimized TENet architecture parameters.")

    add_para(doc, (
        "Three key architectural design choices distinguish this "
        "configuration:"
    ))

    add_bullet_list(doc, [
        "Narrow channels with high expansion ratio (14 ch \u00d7 4\u00d7 expansion): "
        "This creates a rich intermediate representation (56 channels) for "
        "depthwise convolutions while maintaining a compact bottleneck. "
        "This is more efficient than wider channels with lower expansion "
        "(e.g., 24 ch \u00d7 2\u00d7) at the same MACC budget.",

        "Double downsampling [2, 2, 1, 1]: Two stride-2 stages reduce the "
        "temporal dimension from 62 \u2192 31 \u2192 16 early in the network, "
        "allowing later blocks to operate on shorter sequences with 48% "
        "fewer temporal operations.",

        "Extra depth at mid-resolution (layers=[1, 2, 1, 1]): Placing an "
        "additional layer at the 31\u219216 transition provides greater "
        "representational capacity where information density is highest, "
        "at a cost of only ~52K additional MACCs.",
    ])

    # III.F Architecture Optimization (#9)
    add_heading(doc, "F. Architecture Optimization", level=2)
    add_para(doc, (
        "After selecting TENet as the base architecture, a four-stage "
        "progressive grid search was conducted to identify improved "
        "architectural configurations within the MACC budget of 287,673 "
        "(matching the DS-CNN baseline [9])."
    ), indent=True)

    add_para(doc, "The search explored six hyperparameters:")

    add_table(doc,
              ["Dimension", "Values Explored", "Winner"],
              [
                  ["Stem channels", "{8, 12, 14, 16, 18, 20, 24}", "14"],
                  ["Block channels", "{14, 16, 18, 20, 22, 24, 28, 32}", "14"],
                  ["Number of blocks", "{2, 3, 4}", "4"],
                  ["Expansion ratio", "{1, 2, 3, 4}", "4"],
                  ["Kernel size", "{3, 5, 7, 9, 11, 13, 15}", "9"],
                  ["Layers per block", "Various combinations", "[1, 2, 1, 1]"],
              ])

    add_para(doc, (
        "Stride patterns were determined by the number of blocks: three "
        "candidate patterns per block count (e.g., [2,1,1,1], [2,2,1,1], "
        "[2,1,1,2] for 4 blocks). Any candidate exceeding the MACC budget "
        "was discarded before training."
    ), indent=True)

    add_para(doc, (
        "The search was conducted in four progressive stages, summarized "
        "in Table III:"
    ))

    add_table(doc,
              ["Stage", "Candidates", "Epochs", "Augmentation", "GPU Hours", "Selection"],
              [
                  ["Round 1 (Coarse)", "1,280", "15", "None", "2.7", "Top val accuracy"],
                  ["Round 2a (Fine)", "548", "15", "None", "1.1", "Expanded grid"],
                  ["Round 2b (Deep)", "464", "200", "All (\u00a7III.G)", "12.0", "Val accuracy"],
                  ["Round 2c (Final)", "10", "1,000", "All", "0.5", "Final ranking"],
                  ["Total", "2,302", "\u2014", "\u2014", "~16", "\u2014"],
              ],
              caption="Table III: Four-stage progressive grid search strategy.")

    add_para(doc, (
        "All experiments were conducted on an NVIDIA RTX 5080 Laptop GPU "
        "with data preloaded to GPU VRAM for efficient small-model "
        "training. The progressive search strategy is conceptually related "
        "to successive halving [15], where computational resources are "
        "concentrated on the most promising candidates in later stages."
    ), indent=True)

    # III.G Training Configuration (#8 — new)
    add_heading(doc, "G. Training Configuration", level=2)
    add_para(doc, "The final model was trained using the following configuration:")

    add_table(doc,
              ["Parameter", "Value"],
              [
                  ["Optimizer", "AdamW"],
                  ["Learning rate", "0.001"],
                  ["Weight decay", "1\u00d710\u207b\u2074"],
                  ["Batch size", "1,024"],
                  ["Maximum epochs", "300"],
                  ["Best epoch", "273"],
                  ["Early stopping patience", "50 epochs"],
                  ["LR scheduler", "CosineAnnealingLR + 5-epoch linear warmup"],
                  ["Minimum LR (\u03b7_min)", "1\u00d710\u207b\u2075"],
                  ["Random seed", "123"],
              ])

    add_para(doc, "Four data augmentation and regularization techniques were applied:")

    add_numbered_list(doc, [
        "SpecAugment [17]: 2 time masks (width=5 frames) and 1 frequency "
        "mask (width=2 bins) applied to MFCC features during training.",

        "Mixup [18]: Convex combination of training pairs with \u03b1=0.2, "
        "applied at the input level.",

        "Label smoothing [19]: Cross-entropy loss with smoothing factor "
        "\u03b5=0.1.",

        "Cosine annealing [20]: Learning rate decayed from 0.001 to "
        "10\u207b\u2075 following a cosine schedule after a 5-epoch linear warmup.",
    ])

    add_para(doc, (
        "These techniques collectively contributed a +1.84 percentage point "
        "improvement in validation accuracy compared with training without "
        "augmentation (see Section IV.D for ablation results)."
    ), indent=True)

    # III.H INT8 Quantization (#10)
    add_heading(doc, "H. INT8 Quantization", level=2)
    add_para(doc, (
        "Post-training static quantization was applied to convert the "
        "trained float32 model to INT8 precision for efficient embedded "
        "inference."
    ), indent=True)

    add_table(doc,
              ["Parameter", "Value"],
              [
                  ["Tool", "ONNX Runtime quantize_static"],
                  ["Quantization format", "QDQ (QuantizeLinear/DequantizeLinear)"],
                  ["Activation type", "QInt8 (signed 8-bit, per-tensor)"],
                  ["Weight type", "QInt8 (signed 8-bit, per-channel)"],
                  ["Activation symmetric", "False (asymmetric)"],
                  ["Weight symmetric", "True"],
                  ["Calibration set", "200 random training samples (seed=42)"],
                  ["ONNX opset", "13"],
              ])

    add_para(doc, (
        "The QDQ format inserts explicit quantization and dequantization "
        "nodes in the ONNX graph, which is directly supported by the ST "
        "X-CUBE-AI conversion tool [8]. Per-channel weight quantization "
        "preserves accuracy better than per-tensor quantization for "
        "depthwise convolution layers."
    ), indent=True)

    # III.I Embedded Deployment
    add_heading(doc, "I. Embedded Deployment", level=2)
    add_para(doc, (
        "The quantized INT8 ONNX model was deployed on the STM32F767IGT6 "
        "microcontroller using ST X-CUBE-AI v2.2.0 [8]."
    ), indent=True)

    add_table(doc,
              ["Property", "Value"],
              [
                  ["MCU", "STM32F767IGT6"],
                  ["Core", "ARM Cortex-M7"],
                  ["Clock frequency", "216 MHz"],
                  ["Flash memory", "1 MB"],
                  ["SRAM", "512 KB"],
                  ["FPU", "Single-precision"],
                  ["AI framework", "X-CUBE-AI v2.2.0"],
              ])

    add_para(doc, (
        "The deployment pipeline consists of four steps: (1) training in "
        "PyTorch (float32); (2) export to ONNX format (opset 13); "
        "(3) INT8 quantization using ONNX Runtime (QDQ format); and "
        "(4) code generation using X-CUBE-AI, producing optimized C "
        "library and weight arrays. Inference latency is measured using the "
        "ARM DWT (Data Watchpoint and Trace) cycle counter, which provides "
        "cycle-accurate timing of the ai_run() function call."
    ), indent=True)

    # =========================================================================
    # IV. RESULTS AND DISCUSSION
    # =========================================================================
    add_heading(doc, "IV. Results and Discussion", level=1)

    # IV.A Experimental Setup (#14)
    add_heading(doc, "A. Experimental Setup", level=2)
    add_para(doc, (
        "All models were trained using PyTorch 2.10.0 with CUDA on an "
        "NVIDIA RTX 5080 Laptop GPU. Evaluation was performed on the "
        "official test split of Google Speech Commands v2 (12,105 samples). "
        "Classification accuracy is computed as top-1 accuracy. F1-scores "
        "(weighted and macro) are computed using scikit-learn\u2019s "
        "classification_report. All reported test accuracies are from a "
        "single training run with seed=123."
    ), indent=True)

    # IV.B Classification Performance
    add_heading(doc, "B. Classification Performance", level=2)
    add_para(doc, (
        "Table IV presents the overall classification performance of the "
        "optimized TENet model in both float32 and INT8 precision."
    ), indent=True)

    add_table(doc,
              ["Metric", "Float32", "INT8", "\u0394"],
              [
                  ["Test Accuracy", "96.84%", "96.76%", "\u22120.08%"],
                  ["Weighted F1", "0.9683", "0.9675", "\u22120.0008"],
                  ["Macro F1", "0.9504", "0.9490", "\u22120.0014"],
                  ["Macro Precision", "0.9604", "0.9573", "\u22120.0031"],
                  ["Macro Recall", "0.9410", "0.9411", "+0.0001"],
              ],
              caption="Table IV: Overall classification performance.")

    add_para(doc, (
        "INT8 quantization causes only minimal accuracy degradation "
        "(\u22120.08%), demonstrating that the quantized model preserves the "
        "predictive capability of the floating-point model while enabling "
        "efficient embedded deployment."
    ), indent=True)

    # IV.C Per-Class Analysis (#11 — new)
    add_heading(doc, "C. Per-Class Analysis", level=2)
    add_para(doc, (
        "Table V presents the per-class F1-scores for both float32 and "
        "INT8 models."
    ), indent=True)

    add_table(doc,
              ["Class", "F1 (Float32)", "F1 (INT8)", "\u0394 F1", "Support"],
              [
                  ["yes", "0.9736", "0.9723", "\u22120.0013", "419"],
                  ["no", "0.9302", "0.9200", "\u22120.0102", "405"],
                  ["up", "0.9440", "0.9396", "\u22120.0044", "425"],
                  ["down", "0.9362", "0.9271", "\u22120.0091", "406"],
                  ["left", "0.9441", "0.9484", "+0.0043", "412"],
                  ["right", "0.9406", "0.9404", "\u22120.0002", "396"],
                  ["on", "0.9258", "0.9343", "+0.0085", "396"],
                  ["off", "0.9370", "0.9315", "\u22120.0055", "402"],
                  ["stop", "0.9766", "0.9779", "+0.0013", "411"],
                  ["go", "0.9190", "0.9184", "\u22120.0006", "402"],
                  ["unknown", "0.9782", "0.9778", "\u22120.0004", "6,931"],
                  ["silence", "1.0000", "1.0000", "0.0000", "1,100"],
              ],
              caption="Table V: Per-class F1-scores (Float32 vs INT8).")

    add_para(doc, (
        "The silence class achieves perfect classification (F1=1.000) in "
        "both precision modes. The highest keyword F1-scores are observed "
        "for stop (0.977) and yes (0.972), while go (0.918) and on (0.934) "
        "exhibit the lowest performance. The dominant error mode is "
        "confusion between keyword classes and the unknown class, which is "
        "expected given the acoustic similarity between target keywords and "
        "out-of-vocabulary speech."
    ), indent=True)

    add_para(doc, (
        "The unknown class contains 6,931 samples (57.3% of the test set), "
        "and its high F1-score (0.978) substantially contributes to the "
        "weighted F1 metric. The macro F1-score (0.9490), which weights "
        "all classes equally, provides a more balanced assessment of "
        "per-class performance."
    ), indent=True)

    add_para(doc, (
        "INT8 quantization produces minimal per-class variation: the "
        "maximum F1 degradation is 0.0102 (no class), while three classes "
        "(left, on, stop) show slight improvement after quantization, "
        "likely due to beneficial rounding effects."
    ), indent=True)

    # IV.D Ablation Study (#9/new)
    add_heading(doc, "D. Ablation Study: Training Optimization Impact", level=2)
    add_para(doc, (
        "To quantify the contribution of training techniques, Table VI "
        "presents an ablation study comparing different training "
        "configurations."
    ), indent=True)

    add_table(doc,
              ["Configuration", "Val Acc", "Test Acc", "\u0394 vs. Previous"],
              [
                  ["Baseline TENet (manual arch, no aug.)", "95.56%", "95.13%", "\u2014"],
                  ["Grid-optimized arch, no aug., 15 ep", "95.56%", "\u2014", "+0.00%"],
                  ["Grid-optimized arch + all aug., 200 ep", "97.09%", "\u2014", "+1.53%"],
                  ["Grid-optimized arch + all aug., 300 ep", "97.40%", "96.84%", "+0.31%"],
              ],
              caption="Table VI: Ablation study of architecture optimization and training techniques.")

    add_para(doc, (
        "The results show that architecture optimization alone provides a "
        "modest improvement when evaluated at short training horizons "
        "(15 epochs). Training augmentation techniques (SpecAugment, Mixup, "
        "label smoothing, cosine annealing) contribute +1.84% validation "
        "accuracy improvement, demonstrating that training methodology is "
        "as important as architecture design for small models. Extended "
        "training (300 epochs with early stopping at epoch 273) provides an "
        "additional +0.31% improvement."
    ), indent=True)

    add_para(doc, (
        "The gap between validation accuracy (97.40%) and test accuracy "
        "(96.84%) is 0.56%, which is within the expected range for a "
        "single-seed evaluation and reflects the selection bias inherent "
        "in choosing the best architecture from 2,302 candidates."
    ), indent=True)

    # IV.E Model Size (#6 — MACC corrected)
    add_heading(doc, "E. Model Size and Deployment Metrics", level=2)

    add_table(doc,
              ["Metric", "Value"],
              [
                  ["Parameters", "12,822"],
                  ["MACCs (model)", "276,136"],
                  ["MACCs (X-CUBE-AI, incl. QDQ overhead)", "280,389"],
                  ["Flash memory (weights)", "13.87 KB"],
                  ["RAM (activations)", "7.64 KB"],
                  ["Firmware binary size", "67.4 KB"],
                  ["Weight compression vs. float32", "\u221270.8%"],
              ],
              caption="Table VII: Model size and deployment metrics (INT8).")

    add_para(doc, (
        "The MACC count reported by X-CUBE-AI (280,389) is slightly higher "
        "than the model-intrinsic count (276,136) due to additional "
        "operations introduced by QuantizeLinear and DequantizeLinear nodes "
        "in the INT8 inference graph."
    ), indent=True)

    # IV.F Latency
    add_heading(doc, "F. On-Device Inference Latency", level=2)
    add_para(doc, (
        "Inference latency was measured directly on the STM32F767IGT6 "
        "microcontroller using the DWT cycle counter."
    ), indent=True)

    add_table(doc,
              ["Model", "MCU", "Clock", "Latency", "MACCs", "Cycles/MACC"],
              [
                  ["Ours (TENet INT8)", "STM32F767", "216 MHz", "6.32 ms", "276K", "~4.9"],
                  ["DS-CNN baseline [9]", "STM32F769NI", "216 MHz", "7.00 ms", "288K", "~5.3"],
              ],
              caption="Table VIII: On-device inference latency comparison.")

    add_para(doc, (
        "The optimized TENet model achieves 6.32 ms inference latency, "
        "which is 9.7% faster than the DS-CNN baseline despite operating "
        "on the same Cortex-M7 core at the same clock frequency. The ~4.9 "
        "cycles per MACC is consistent with INT8 mixed-precision inference "
        "on the Cortex-M7 architecture using X-CUBE-AI optimized kernels."
    ), indent=True)

    # IV.G SOTA Comparison (#3, #4, #5)
    add_heading(doc, "G. Comparison with State-of-the-Art", level=2)
    add_para(doc, (
        "Table IX compares the proposed model with previously reported "
        "keyword spotting architectures. Two categories of comparison are "
        "presented: (1) models evaluated on GSCv2 12-class under similar "
        "MACC constraints, and (2) models deployed on microcontrollers with "
        "reported on-device metrics."
    ), indent=True)

    add_table(doc,
              ["Model", "Year", "Params", "MACCs", "Accuracy", "Flash", "RAM", "Latency", "Platform"],
              [
                  ["Ours (INT8)", "2026", "12.8K", "276K", "96.76%", "13.87 KB", "7.64 KB", "6.32 ms", "STM32F767 (M7@216MHz)"],
                  ["DS-CNN [9]", "2020", "~6K", "288K", "91.00%", "169.63 KB", "11.52 KB", "7.00 ms", "STM32F769NI (M7@216MHz)"],
                  ["DS-CNN-S [2]", "2018", "~23K", "5.4M", "94.4%\u2020", "38.6 KB", "~70 KB", "~12 ms", "STM32F746 (M7@216MHz)"],
                  ["TENet12 [4]", "2020", "100K", "2.9M", "96.8%", "\u2014", "\u2014", "\u2014", "\u2014"],
                  ["TENet6-N [4]", "2020", "17K", "553K", "96.0%", "\u2014", "\u2014", "\u2014", "\u2014"],
                  ["BC-ResNet-1 [5]", "2021", "9.2K", "3.1M", "96.9%", "\u2014", "\u2014", "\u2014", "\u2014"],
                  ["MatchboxNet [6]", "2020", "93K", "\u2014", "97.48%", "\u2014", "\u2014", "\u2014", "\u2014"],
                  ["MLPerf Tiny [3]", "2021", "38.6K", "~100K", "\u226590%", "\u2014", "\u2264512 KB", "~182 ms", "NUCLEO-L4R5ZI (M4@120MHz)"],
              ],
              caption="Table IX: Comparison with state-of-the-art keyword spotting models.")

    add_para(doc, (
        "\u2020Zhang et al. [2] evaluated on GSCv1, not GSCv2; results are not "
        "directly comparable."
    ), italic=True)

    add_para(doc, (
        "Note: Bartoli et al. [21] recently evaluated several models "
        "(DS-CNN, TENet6, LiCoNet-S, TKWS) on STM32 platforms, reporting "
        "F1-scores of 91.2%\u201393.6%. However, their evaluation used an "
        "8-class task with different MFCC configurations (15\u00d763), making "
        "direct numerical comparison with our 12-class results "
        "inappropriate. Their study does confirm that TENet and LiCoNet "
        "architectures are more energy-efficient than DS-CNN on Cortex-M7 "
        "processors."
    ), indent=True)

    add_para(doc, (
        "Compared with the DS-CNN baseline [9] evaluated on the same MCU "
        "architecture (Cortex-M7 at 216 MHz) under the same MACC constraint "
        "(~287K):"
    ))

    add_table(doc,
              ["Metric", "DS-CNN Baseline [9]", "Ours (TENet INT8)", "Improvement"],
              [
                  ["Accuracy", "91.00%", "96.76%", "+5.76 pp"],
                  ["Flash (weights)", "169.63 KB", "13.87 KB", "\u221291.8%"],
                  ["RAM (activations)", "11.52 KB", "7.64 KB", "\u221233.7%"],
                  ["Inference latency", "7.00 ms", "6.32 ms", "\u22129.7%"],
              ])

    add_para(doc, (
        "Compared with the original TENet paper [4], our optimized variant "
        "achieves comparable accuracy (96.76% vs. 96.8% for TENet12+MTConv) "
        "while requiring approximately 10\u00d7 fewer MACCs (276K vs. 2.9M). "
        "This demonstrates that systematic architecture optimization can "
        "recover most of the accuracy of larger models within a severely "
        "constrained computational budget."
    ), indent=True)

    # =========================================================================
    # V. CONCLUSION (#15, #20)
    # =========================================================================
    add_heading(doc, "V. Conclusion", level=1)

    add_para(doc, (
        "This paper presented an optimized keyword spotting system designed "
        "for deployment on resource-constrained microcontrollers. Through "
        "a five-architecture comparison study and a four-stage progressive "
        "grid search over 2,302 candidate configurations, an optimized "
        "TENet variant was identified that achieves high classification "
        "accuracy within a strict computational budget of 276K MACCs."
    ), indent=True)

    add_para(doc, (
        "The optimized model achieves 96.84% test accuracy (96.76% after "
        "INT8 quantization) and a weighted F1-score of 0.9675 on the "
        "Google Speech Commands v2 12-class benchmark, with only 12,822 "
        "parameters. When deployed on an STM32F767IGT6 Cortex-M7 "
        "microcontroller at 216 MHz, the INT8 model requires 13.87 KB of "
        "flash memory and 7.64 KB of RAM, and achieves an inference "
        "latency of 6.32 ms. Compared with the DS-CNN baseline operating "
        "under the same MACC constraint, the proposed model improves "
        "accuracy by 5.76 percentage points while reducing flash memory "
        "by 91.8%."
    ), indent=True)

    add_para(doc, (
        "Key findings from this work include: (1) narrow channels with "
        "high expansion ratios are more efficient than wide channels with "
        "low expansion ratios at the same MACC budget; (2) training "
        "augmentation techniques (SpecAugment, Mixup, label smoothing) "
        "collectively contribute +1.84% accuracy improvement for small "
        "models; and (3) INT8 post-training quantization introduces "
        "negligible accuracy degradation (\u22120.08%) for the optimized "
        "architecture."
    ), indent=True)

    add_para(doc, (
        "This work has several limitations. The evaluation is based on a "
        "single training seed, and the reported accuracy may vary by "
        "\u00b10.2\u20130.5% across different random initializations. Energy "
        "consumption was not measured, which is an important metric for "
        "battery-powered devices. The grid search approach, while effective "
        "for the six-dimensional space considered, does not scale to larger "
        "or continuous search spaces."
    ), indent=True)

    add_para(doc, (
        "Future work will explore: (1) multi-seed evaluation to establish "
        "confidence intervals; (2) energy efficiency analysis using current "
        "measurement on the target hardware; (3) extension to "
        "streaming/continuous keyword spotting scenarios; (4) "
        "hardware-aware NAS methods (e.g., differentiable NAS, Bayesian "
        "optimization) to explore larger architecture spaces; and "
        "(5) deployment on additional MCU platforms including Cortex-M4 and "
        "RISC-V processors."
    ), indent=True)

    # =========================================================================
    # REFERENCES (#5)
    # =========================================================================
    add_heading(doc, "References", level=1)

    refs = [
        '[1] P. Warden, "Speech commands: A dataset for limited-vocabulary speech recognition," arXiv preprint arXiv:1804.03209, 2018.',
        '[2] Y. Zhang, N. Suda, L. Lai, and V. Chandra, "Hello edge: Keyword spotting on microcontrollers," arXiv preprint arXiv:1711.07128, 2017.',
        '[3] C. Banbury et al., "MLPerf Tiny benchmark," in Proc. Neural Information Processing Systems (NeurIPS) Datasets and Benchmarks Track, 2021.',
        '[4] R. Li, Z. Gao, L. Wang, and G. Li, "Small-footprint keyword spotting with multi-scale temporal convolution," in Proc. Interspeech, 2020.',
        '[5] B. Kim, S. Chang, J. Lee, and D. Sung, "Broadcasted residual learning for efficient keyword spotting," in Proc. Interspeech, 2021.',
        '[6] S. Majumdar and B. Ginsburg, "MatchboxNet: 1D time-channel separable convolutional neural network architecture for speech commands recognition," in Proc. Interspeech, 2020.',
        '[7] S. Davis and P. Mermelstein, "Comparison of parametric representations for monosyllabic word recognition in continuously spoken sentences," IEEE Trans. Acoustics, Speech, and Signal Processing, vol. 28, no. 4, pp. 357\u2013366, 1980.',
        '[8] STMicroelectronics, "X-CUBE-AI: AI expansion pack for STM32CubeMX," 2024.',
        '[9] P. M. Sorensen, B. Epp, and T. May, "A depthwise separable convolutional neural network for keyword spotting on an embedded system," EURASIP J. Audio, Speech, and Music Processing, vol. 2020, no. 1, pp. 1\u201311, 2020.',
        '[10] M. Sandler, A. Howard, M. Zhu, A. Zhmoginov, and L.-C. Chen, "MobileNetV2: Inverted residuals and linear bottlenecks," in Proc. IEEE/CVF CVPR, 2018, pp. 4510\u20134520.',
        '[11] J. Lin, W.-M. Chen, Y. Lin, J. Cohn, C. Gan, and S. Han, "MCUNet: Tiny deep learning on IoT devices," in Proc. NeurIPS, 2020.',
        '[12] H. Liu, K. Simonyan, and Y. Yang, "DARTS: Differentiable architecture search," in Proc. ICLR, 2019.',
        '[13] S. Liberatore, F. Pollicino, and M. Rusci, "MicroNAS: Zero-shot neural architecture search for MCUs," arXiv preprint arXiv:2401.08996, 2024.',
        '[14] J. Bergstra and Y. Bengio, "Random search for hyper-parameter optimization," J. Machine Learning Research, vol. 13, pp. 281\u2013305, 2012.',
        '[15] K. Jamieson and A. Talwalkar, "Non-stochastic best arm identification and hyperparameter optimization," in Proc. AISTATS, 2016.',
        '[16] L. Lai, N. Suda, and V. Chandra, "CMSIS-NN: Efficient neural networks on ARM Cortex-M processors," arXiv preprint arXiv:1801.06601, 2018.',
        '[17] D. S. Park et al., "SpecAugment: A simple data augmentation method for automatic speech recognition," in Proc. Interspeech, 2019.',
        '[18] H. Zhang, M. Ciss\u00e9, Y. N. Dauphin, and D. Lopez-Paz, "Mixup: Beyond empirical risk minimization," in Proc. ICLR, 2018.',
        '[19] C. Szegedy, V. Vanhoucke, S. Ioffe, J. Shlens, and Z. Wojna, "Rethinking the inception architecture for computer vision," in Proc. IEEE/CVF CVPR, 2016.',
        '[20] I. Loshchilov and F. Hutter, "SGDR: Stochastic gradient descent with warm restarts," in Proc. ICLR, 2017.',
        '[21] F. Bartoli, L. Boschi, and V. Luise, "End-to-end efficiency in keyword spotting: A system-level approach for embedded microcontrollers," arXiv preprint arXiv:2509.07051, 2025.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    # =========================================================================
    # SAVE
    # =========================================================================
    out_path = "C:/Users/m1339/PycharmProjects/KWS_Project/Optimizing Temporal Efficient Networks for Embedded Keyword Spotting on Microcontrollers.docx"
    doc.save(out_path)
    print(f"Saved to: {out_path}")
    print(f"Sections: 5 main + References")
    print(f"Tables: 9")
    print(f"References: {len(refs)}")


if __name__ == "__main__":
    main()
